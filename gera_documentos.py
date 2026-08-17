#!/usr/bin/env python3
"""
gera_documentos.py — Converte o README.md em versões .txt e .docx
para facilitar o envio na plataforma da faculdade (Unidade IV).
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

BASE = Path("/home/elissandrareis/Área de Trabalho/ATIVIDADE PRÁTICA – UNIDADE IV")


# ---------------------------------------------------------------------------
# Versão .txt (texto puro, sem formatação markdown)
# ---------------------------------------------------------------------------
def gerar_txt():
    md = (BASE / "README.md").read_text(encoding="utf-8")

    # Remove blocos de código, mantém o conteúdo
    txt = md.replace("```python", "```").replace("```bash", "```").replace("```", "")

    # Remove marcadores markdown (**negrito**, `código`, _itálico_)
    txt = re.sub(r"\*\*(.+?)\*\*", r"\1", txt)   # **negrito**
    txt = re.sub(r"(?<![\w])\*([^*]+)\*(?![\w])", r"\1", txt)  # *itálico*
    txt = re.sub(r"`([^`]+)`", r"\1", txt)           # `código`
    txt = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", txt)  # [texto](url)

    # Tabelas markdown -> linhas simples
    linhas = []
    for linha in txt.splitlines():
        if linha.strip().startswith("|"):
            celulas = [c.strip() for c in linha.strip().strip("|").split("|")]
            if set("".join(celulas).replace("-", "").replace(":", "")) == set():
                continue  # linha separadora da tabela
            linhas.append(" | ".join(celulas))
        else:
            linhas.append(linha)
    txt = "\n".join(linhas)

    # Títulos markdown -> letras maiúsculas
    for nivel in ("#### ", "### ", "## ", "# "):
        txt = txt.replace(nivel, "")

    out = BASE / "EXPLICACAO.txt"
    out.write_text(txt, encoding="utf-8")
    print(f"gerado: {out.name}")


# ---------------------------------------------------------------------------
# Versão .docx (Word, formatado)
# ---------------------------------------------------------------------------
def gerar_docx():
    doc = Document()

    # Estilo base
    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)

    AZUL = RGBColor(0x1F, 0x4E, 0x79)

    def titulo(texto, nivel=0):
        h = doc.add_heading(texto, level=nivel)
        for run in h.runs:
            run.font.color.rgb = AZUL
        return h

    def paragrafo(texto, negrito=False, italico=False):
        p = doc.add_paragraph()
        r = p.add_run(texto)
        r.bold = negrito
        r.italic = italico
        return p

    def codigo(texto):
        p = doc.add_paragraph()
        r = p.add_run(texto)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        return p

    def tabela(cabecalho, linhas):
        t = doc.add_table(rows=1, cols=len(cabecalho))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, c in enumerate(cabecalho):
            t.rows[0].cells[i].text = c
            for r in t.rows[0].cells[i].paragraphs[0].runs:
                r.bold = True
        for linha in linhas:
            celulas = t.add_row().cells
            for i, valor in enumerate(linha):
                celulas[i].text = str(valor)
        return t

    # ---------------- CAPA ----------------
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("API REST de Jogos")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = AZUL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Atividade Prática — Unidade IV\nFastAPI • JSON • Persistência no Cloud Firestore\nFirebase Admin SDK")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    doc.add_page_break()

    # ---------------- 1. TEMA ----------------
    titulo("1. Tema escolhido", 1)
    paragrafo(
        "A API gerencia um catálogo de jogos (videogames) — tema diferente do exemplo "
        "utilizado pelo professor em sala de aula (alunos e professores). É a mesma API "
        "desenvolvida na Unidade III, evoluída para persistir os dados no Cloud Firestore."
    )
    paragrafo(
        "Cada registro representa um jogo e possui os mesmos 5 atributos da Unidade III:",
        negrito=True
    )
    tabela(
        ["Atributo", "Tipo", "Descrição"],
        [
            ["id", "int", "Identificador único e obrigatório"],
            ["titulo", "str", "Nome do jogo"],
            ["genero", "str", "Gênero (Aventura, Ação, RPG, Esporte, ...)"],
            ["plataforma", "str", "Plataforma (PC, PlayStation 5, Xbox, ...)"],
            ["ano_lancamento", "int", "Ano de lançamento"],
        ],
    )

    # ---------------- 2. O QUE MUDOU ----------------
    titulo("2. O que mudou: lista em memória -> Cloud Firestore", 1)
    paragrafo(
        "Na Unidade III, os registros eram mantidos em uma lista Python (list) existente "
        "apenas na memória da aplicação. Por isso, ao encerrar ou reiniciar o servidor, "
        "todos os dados eram perdidos."
    )
    paragrafo(
        "Nesta Unidade IV, a lista foi substituída pelo Cloud Firestore, banco de dados "
        "NoSQL do Firebase. As mudanças no código foram:"
    )
    paragrafo("• A variável global jogos = [] foi removida;", negrito=False)
    paragrafo(
        "• Foi criada uma coleção \"jogos\" no Firestore, e cada registro passou a ser um "
        "DOCUMENTO dessa coleção. O identificador único (id) virou o id do documento "
        "(ex.: o jogo com id 777 é o documento \"777\"), o que garante a unicidade e "
        "permite leitura pontual por id;"
    )
    paragrafo(
        "• POST /jogos agora usa doc_ref.create() para gravar o documento: se o id já "
        "existir, o Firestore lança Conflict e a API responde 409 (mesmo comportamento "
        "da Unidade III, agora garantido pelo banco de dados);"
    )
    paragrafo(
        "• GET /jogos lê todos os documentos da coleção com colecao.stream(); o filtro "
        "opcional ?genero=... continua ignorando maiúsculas/minúsculas (comportamento "
        "preservado da Unidade III);"
    )
    paragrafo(
        "• GET /jogos/{id} faz uma leitura pontual do documento no Firestore e responde "
        "404 Not Found quando o documento não existe."
    )
    paragrafo(
        "Resultado: os registros criados pela API ficam gravados no Firestore e "
        "continuam disponíveis mesmo após o encerramento ou a reinicialização da "
        "aplicação.", negrito=True
    )

    # ---------------- 3. CONFIGURAÇÃO ----------------
    titulo("3. Configuração necessária (credenciais)", 1)
    paragrafo(
        "O mecanismo de integração é o Firebase Admin SDK (firebase-admin), "
        "conforme a documentação oficial indicada no material da disciplina "
        "(firebase.google.com/docs/firestore/manage-data/add-data). Nenhuma "
        "chave privada é incluída na entrega; a credencial deve ser fornecida "
        "de uma das formas abaixo (o código verifica na ordem: arquivo -> "
        "variável de ambiente -> emulador):"
    )
    paragrafo(
        "1) Mecanismo indicado no material: no Console do Firebase "
        "(console.firebase.google.com), crie o projeto, habilite o Cloud "
        "Firestore e, em Configurações do projeto -> Contas de serviço, clique "
        "em Gerar nova chave privada. Salve o JSON baixado como "
        "serviceAccountKey.json na pasta do projeto (credentials.Certificate):"
    )
    codigo("uvicorn main:app --reload")
    paragrafo(
        "Ou informe a chave por variável de ambiente (credenciais padrão do Google):"
    )
    codigo("export GOOGLE_APPLICATION_CREDENTIALS=/caminho/para/chave.json\nuvicorn main:app --reload")
    paragrafo(
        "2) Emulador local (testes): defina FIRESTORE_EMULATOR_HOST apontando para "
        "o emulador do Firestore, por exemplo:"
    )
    codigo("export FIRESTORE_EMULATOR_HOST=127.0.0.1:8081\nuvicorn main:app --reload")

    # ---------------- 4. ROTAS ----------------
    titulo("4. Rotas, métodos HTTP e códigos de resposta (mantidos)", 1)
    tabela(
        ["Método", "Rota", "Função", "Códigos de resposta"],
        [
            ["POST", "/jogos", "Criar um novo jogo", "201 Created / 409 Conflict / 422"],
            ["GET", "/jogos", "Listar todos os jogos", "200 OK"],
            ["GET", "/jogos?genero=RPG", "Listar jogos de um gênero", "200 OK"],
            ["GET", "/jogos/{jogo_id}", "Consultar um jogo pelo id", "200 OK / 404 Not Found"],
        ],
    )

    doc.add_paragraph()
    titulo("4.1 POST /jogos — criar um novo registro", 2)
    paragrafo(
        "O documento é criado com o id do jogo (ex.: id=7 vira o documento \"7\"). "
        "Como o Firestore não permite dois documentos com o mesmo id na mesma coleção, "
        "o identificador único é garantido: se o id já existir, o método create() "
        "lança AlreadyExists e a API responde 409 Conflict. O 422 é gerado "
        "automaticamente pelo FastAPI para JSON fora do esquema."
    )

    titulo("4.2 GET /jogos — listar registros (com filtro opcional)", 2)
    paragrafo(
        "Lista todos os jogos persistidos no Firestore e retorna 200 OK com um array "
        "JSON. A rota possui o parâmetro de consulta opcional genero "
        "(/jogos?genero=RPG):"
    )
    paragrafo("• Parâmetro não informado -> retorna a lista completa de jogos.")
    paragrafo(
        "• Parâmetro informado -> retorna apenas os jogos do gênero correspondente. A "
        "comparação ignora maiúsculas/minúsculas (ex.: ?genero=rpg encontra \"RPG\")."
    )

    titulo("4.3 GET /jogos/{jogo_id} — consultar por ID", 2)
    paragrafo(
        "Localiza o documento do jogo diretamente pelo id no Firestore. Retorna 200 OK "
        "com o registro encontrado ou 404 Not Found quando o id não existe."
    )

    # ---------------- 5. PERSISTÊNCIA ----------------
    titulo("5. Comprovando a persistência", 1)
    paragrafo(
        "Procedimento realizado para comprovar que os dados ficam no Firestore mesmo "
        "após a reinicialização da aplicação:"
    )
    paragrafo("1. Criou-se um registro pela API (POST /jogos com id 777 -> 201 Created);")
    paragrafo(
        "2. O documento apareceu na coleção jogos do Firestore com todos os atributos;"
    )
    paragrafo("3. A aplicação foi encerrada e iniciada novamente (reinicialização);")
    paragrafo(
        "4. Consultou-se o mesmo registro (GET /jogos/777 -> 200 OK): o registro "
        "continuou disponível, pois foi lido do Firestore e não da memória.",
        negrito=True,
    )
    paragrafo(
        "O script demonstra_persistencia.py automatiza esse procedimento."
    )

    # ---------------- 6. CAPTURAS ----------------
    titulo("6. Testes realizados em /docs (capturas de tela)", 1)
    tabela(
        ["Captura", "Teste realizado", "Resultado"],
        [
            ["01_docs_visao_geral.png", "Visão geral das rotas no Swagger UI", "—"],
            ["02_post_jogos_201_created.png", "POST /jogos (criar jogo)", "201 Created"],
            ["03_get_jogos_lista_200.png", "GET /jogos (listar todos)", "200 OK"],
            ["04_get_jogos_filtro_genero.png", "GET /jogos?genero=RPG (filtro)", "200 OK"],
            ["05_get_jogos_id_200.png", "GET /jogos/777 (consulta por id)", "200 OK"],
            ["06_get_jogos_id_404_not_found.png", "GET /jogos/999 (id inexistente)", "404 Not Found"],
            ["07_firestore_documento.png", "Documento na coleção jogos do Firestore", "—"],
            ["08_persistencia_apos_reinicio.png", "Consulta após reiniciar a aplicação", "200 OK"],
        ],
    )
    doc.add_paragraph()
    paragrafo(
        "A captura 07 foi feita na UI do emulador do Firestore; com um projeto real, o "
        "documento é visto no Console do Firebase (Firestore Database -> coleção jogos).",
        italico=True,
    )

    # ---------------- 7. ARQUIVOS ----------------
    titulo("7. Arquivos da entrega", 1)
    tabela(
        ["Arquivo", "Descrição"],
        [
            ["main.py", "Aplicação FastAPI integrada ao Cloud Firestore (Firebase Admin SDK)"],
            ["requirements.txt", "Dependências (fastapi, uvicorn, firebase-admin)"],
            ["README.md / EXPLICACAO.txt / EXPLICACAO.docx", "Texto explicativo"],
            ["capturas/", "Capturas de tela dos testes em /docs e da persistência"],
            ["teste_api.py", "Script com 22 testes automáticos"],
            ["demonstra_persistencia.py", "Script que comprova a persistência"],
            ["capturar_docs.py", "Script que reproduz as capturas de tela"],
        ],
    )
    paragrafo(
        "Observação: o arquivo de credenciais (serviceAccountKey.json) NÃO é enviado — "
        "apenas a explicação de como a configuração deve ser fornecida.",
        italico=True,
    )

    out = BASE / "EXPLICACAO.docx"
    doc.save(out)
    print(f"gerado: {out.name}")


if __name__ == "__main__":
    gerar_txt()
    gerar_docx()
